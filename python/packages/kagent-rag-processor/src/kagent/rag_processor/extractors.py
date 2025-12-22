"""Text extraction from various document types."""

import csv
import io
import json
import logging
from pathlib import Path

import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


def extract_text(content: bytes, filename: str) -> str:
    """Extract text from document content based on file extension.

    Args:
        content: Raw file content as bytes
        filename: Original filename with extension

    Returns:
        Extracted text content

    Raises:
        ValueError: If file type is not supported
    """
    ext = Path(filename).suffix.lower()

    extractors = {
        ".txt": _extract_txt,
        ".md": _extract_txt,
        ".json": _extract_json,
        ".csv": _extract_csv,
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
    }

    extractor = extractors.get(ext)
    if not extractor:
        raise ValueError(f"Unsupported file type: {ext}")

    return extractor(content)


def _extract_txt(content: bytes) -> str:
    """Extract text from plain text or markdown files."""
    return content.decode("utf-8", errors="replace")


def _extract_json(content: bytes) -> str:
    """Extract text from JSON files."""
    data = json.loads(content.decode("utf-8"))
    return json.dumps(data, indent=2)


def _extract_csv(content: bytes) -> str:
    """Extract text from CSV files."""
    text_io = io.StringIO(content.decode("utf-8", errors="replace"))
    reader = csv.reader(text_io)
    lines = []
    for row in reader:
        lines.append(" | ".join(row))
    return "\n".join(lines)


def _extract_pdf(content: bytes) -> str:
    """Extract text from PDF files."""
    text_parts = []
    pdf_io = io.BytesIO(content)

    with pdfplumber.open(pdf_io) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    return "\n\n".join(text_parts)


def _extract_docx(content: bytes) -> str:
    """Extract text from DOCX files."""
    docx_io = io.BytesIO(content)
    doc = Document(docx_io)

    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                text_parts.append(row_text)

    return "\n\n".join(text_parts)
